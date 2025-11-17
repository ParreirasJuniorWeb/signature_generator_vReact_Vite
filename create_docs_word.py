"""
Script para converter documentação da pasta docs/ para Word (.docx)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def read_file(filepath):
    """Lê o conteúdo de um arquivo"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        return f'Arquivo não encontrado: {filepath}'
    except Exception as e:
        return f'Erro ao ler arquivo: {e}'

def add_formatted_content(doc, content):
    """Adiciona conteúdo formatado ao documento"""
    lines = content.split('\n')
    in_code_block = False
    code_lines = []
    
    for line in lines:
        # Detectar blocos de código
        if line.strip().startswith('```'):
            if in_code_block:
                # Fim do bloco de código
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph(code_text)
                    p.style = 'Intense Quote'
                code_lines = []
                in_code_block = False
            else:
                # Início do bloco de código
                in_code_block = True
            continue
        
        if in_code_block:
            code_lines.append(line)
            continue
        
        line_stripped = line.strip()
        
        if not line_stripped:
            doc.add_paragraph()
            continue
            
        # Títulos
        if line_stripped.startswith('# '):
            heading = doc.add_heading(line_stripped[2:], level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line_stripped.startswith('## '):
            doc.add_heading(line_stripped[3:], level=2)
        elif line_stripped.startswith('### '):
            doc.add_heading(line_stripped[4:], level=3)
        elif line_stripped.startswith('#### '):
            doc.add_heading(line_stripped[5:], level=4)
        
        # Listas com checkbox
        elif line_stripped.startswith('- [x] ') or line_stripped.startswith('- [ ] '):
            checkbox = '☑' if '[x]' in line_stripped else '☐'
            text = line_stripped.replace('- [x] ', '').replace('- [ ] ', '')
            doc.add_paragraph(f'{checkbox} {text}', style='List Bullet')
        
        # Listas normais
        elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
            doc.add_paragraph(line_stripped[2:], style='List Bullet')
        elif any(line_stripped.startswith(f'{i}. ') for i in range(1, 10)):
            # Encontrar onde termina o número
            dot_pos = line_stripped.find('. ')
            if dot_pos > 0:
                doc.add_paragraph(line_stripped[dot_pos+2:], style='List Number')
        
        # Texto normal
        else:
            doc.add_paragraph(line_stripped)

def create_document(title, input_file, output_file):
    """Cria um documento Word a partir de um arquivo Markdown"""
    print(f'Criando {output_file}...')
    
    doc = Document()
    
    # Título principal
    main_title = doc.add_heading(title, 0)
    main_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Adicionar data
    date_para = doc.add_paragraph('Data: 16 de Novembro de 2025')
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Adicionar status
    status_para = doc.add_paragraph('Status: ✅ Documentação Oficial')
    status_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    status_para.runs[0].bold = True
    
    doc.add_paragraph()
    
    # Ler e adicionar conteúdo
    content = read_file(input_file)
    add_formatted_content(doc, content)
    
    # Salvar
    doc.save(output_file)
    print(f'✅ {output_file} criado com sucesso!')

def main():
    """Função principal"""
    print('🚀 Iniciando geração de documentos Word da pasta docs/...\n')
    
    # Criar pasta docs_word se não existir
    output_dir = 'docs_word'
    os.makedirs(output_dir, exist_ok=True)
    print(f'📁 Usando pasta {output_dir}/\n')
    
    # Lista de documentos para criar
    documents = [
        {
            'title': '📖 API - Documentação da API REST',
            'input': 'docs/API.md',
            'output': f'{output_dir}/API.docx'
        },
        {
            'title': '🏗️ ARCHITECTURE - Arquitetura do Sistema',
            'input': 'docs/ARCHITECTURE.md',
            'output': f'{output_dir}/ARCHITECTURE.docx'
        },
        {
            'title': '📝 CHANGELOG - Histórico de Mudanças',
            'input': 'docs/CHANGELOG.md',
            'output': f'{output_dir}/CHANGELOG.docx'
        },
        {
            'title': '🔄 MIGRATION GUIDE - Guia de Migração',
            'input': 'docs/MIGRATION_GUIDE.md',
            'output': f'{output_dir}/MIGRATION_GUIDE.docx'
        },
        {
            'title': '✅ PROJECT COMPLETION - Conclusão do Projeto',
            'input': 'docs/PROJECT_COMPLETION.md',
            'output': f'{output_dir}/PROJECT_COMPLETION.docx'
        },
        {
            'title': '📊 SUMMARY - Resumo Executivo',
            'input': 'docs/SUMMARY.md',
            'output': f'{output_dir}/SUMMARY.docx'
        }
    ]
    
    # Criar cada documento
    success_count = 0
    for doc_info in documents:
        try:
            create_document(
                doc_info['title'],
                doc_info['input'],
                doc_info['output']
            )
            success_count += 1
        except Exception as e:
            print(f'❌ Erro ao criar {doc_info["output"]}: {e}')
    
    print(f'\n🎉 Processo concluído! {success_count}/{len(documents)} documentos criados.')
    print(f'\n📁 Documentos Word salvos em: {os.path.abspath(output_dir)}/')
    print('\n📄 Arquivos criados:')
    for doc_info in documents:
        if os.path.exists(doc_info['output']):
            print(f'  ✅ {doc_info["output"]}')

if __name__ == '__main__':
    main()
