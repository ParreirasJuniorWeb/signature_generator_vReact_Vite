"""
Script para criar um documento Word consolidado com toda a documentação
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def read_file(filepath):
    """Lê o conteúdo de um arquivo"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        return f'Arquivo não encontrado: {filepath}'
    except Exception as e:
        return f'Erro ao ler arquivo: {e}'

def add_page_break(doc):
    """Adiciona quebra de página"""
    doc.add_page_break()

def add_formatted_content(doc, content):
    """Adiciona conteúdo formatado ao documento"""
    lines = content.split('\n')
    in_code_block = False
    code_lines = []
    
    for line in lines:
        # Detectar blocos de código
        if line.strip().startswith('```'):
            if in_code_block:
                # Fim do bloco de código
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph(code_text)
                    p.style = 'Intense Quote'
                code_lines = []
                in_code_block = False
            else:
                # Início do bloco de código
                in_code_block = True
            continue
        
        if in_code_block:
            code_lines.append(line)
            continue
        
        line = line.strip()
        
        if not line:
            doc.add_paragraph()
            continue
            
        # Títulos
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # Listas com checkbox
        elif line.startswith('- [x] ') or line.startswith('- [ ] '):
            checkbox = '☑' if '[x]' in line else '☐'
            text = line.replace('- [x] ', '').replace('- [ ] ', '')
            doc.add_paragraph(f'{checkbox} {text}', style='List Bullet')
        
        # Listas normais
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif any(line.startswith(f'{i}. ') for i in range(1, 10)):
            doc.add_paragraph(line[3:], style='List Number')
        
        # Texto normal
        else:
            doc.add_paragraph(line)

def create_consolidated_document():
    """Cria documento consolidado com toda a documentação"""
    print('🚀 Criando documento consolidado...\n')
    
    doc = Document()
    
    # Capa
    title = doc.add_heading('📚 DOCUMENTAÇÃO COMPLETA', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    subtitle = doc.add_heading('Gerador de Assinaturas de E-mail', level=1)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    info = doc.add_paragraph('Versão: 2.0.0')
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    date = doc.add_paragraph('Data: 16 de Novembro de 2025')
    date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    status = doc.add_paragraph('Status: ✅ APROVADO PARA PRODUÇÃO')
    status.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    status.runs[0].bold = True
    status.runs[0].font.color.rgb = RGBColor(0, 128, 0)
    
    doc.add_page_break()
    
    # Sumário
    doc.add_heading('📋 SUMÁRIO', level=1)
    doc.add_paragraph()
    
    toc_items = [
        '1. README Principal',
        '2. Relatório Final de Testes',
        '3. Checklist Completo de Testes',
        '4. Guia de Correção CORS',
        '5. Guia de Diagnóstico de Conexão',
        '6. Instruções de Teste'
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
        p.runs[0].bold = True
    
    doc.add_page_break()
    
    # Documentos
    documents = [
        {
            'title': '1. 📖 README PRINCIPAL',
            'file': 'README.md'
        },
        {
            'title': '2. 📊 RELATÓRIO FINAL DE TESTES',
            'file': 'FINAL_TESTING_REPORT.md'
        },
        {
            'title': '3. ✅ CHECKLIST COMPLETO DE TESTES',
            'file': 'COMPLETE_TESTING_CHECKLIST.md'
        },
        {
            'title': '4. 🔧 GUIA DE CORREÇÃO CORS',
            'file': 'CORS_FIX_GUIDE.md'
        },
        {
            'title': '5. 🔍 GUIA DE DIAGNÓSTICO DE CONEXÃO',
            'file': 'DEBUG_CONNECTION.md'
        },
        {
            'title': '6. 🧪 INSTRUÇÕES DE TESTE',
            'file': 'TESTING_INSTRUCTIONS.md'
        }
    ]
    
    for doc_info in documents:
        print(f'Adicionando: {doc_info["title"]}')
        
        # Título da seção
        section_title = doc.add_heading(doc_info['title'], level=1)
        section_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()
        
        # Conteúdo
        content = read_file(doc_info['file'])
        add_formatted_content(doc, content)
        
        # Quebra de página (exceto no último)
        if doc_info != documents[-1]:
            doc.add_page_break()
    
    # Salvar
    output_file = 'DOCUMENTACAO_COMPLETA_GERADOR_ASSINATURAS.docx'
    doc.save(output_file)
    print(f'\n✅ {output_file} criado com sucesso!')
    print(f'📁 Localização: {os.path.abspath(output_file)}')
    
    return output_file

if __name__ == '__main__':
    create_consolidated_document()
    print('\n🎉 Documento consolidado criado com sucesso!')
