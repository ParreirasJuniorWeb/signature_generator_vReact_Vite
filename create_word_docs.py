"""
Script para converter documentação Markdown para Word (.docx)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def read_file(filepath):
    """Lê o conteúdo de um arquivo"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        return f'Arquivo não encontrado: {filepath}'
    except Exception as e:
        return f'Erro ao ler arquivo: {e}'

def add_formatted_content(doc, content):
    """Adiciona conteúdo formatado ao documento"""
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            doc.add_paragraph()
            continue
            
        # Títulos
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # Listas
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
            doc.add_paragraph(line[3:], style='List Number')
        
        # Código
        elif line.startswith('```'):
            continue
        
        # Texto normal
        else:
            p = doc.add_paragraph(line)

def create_document(title, input_file, output_file):
    """Cria um documento Word a partir de um arquivo Markdown"""
    print(f'Criando {output_file}...')
    
    doc = Document()
    
    # Título principal
    main_title = doc.add_heading(title, 0)
    main_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Adicionar data
    date_para = doc.add_paragraph('Data: 16 de Novembro de 2025')
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    # Ler e adicionar conteúdo
    content = read_file(input_file)
    add_formatted_content(doc, content)
    
    # Salvar
    doc.save(output_file)
    print(f'✅ {output_file} criado com sucesso!')

def main():
    """Função principal"""
    print('🚀 Iniciando geração de documentos Word...\n')
    
    # Lista de documentos para criar
    documents = [
        {
            'title': '📊 Relatório Final de Testes - Gerador de Assinaturas',
            'input': 'FINAL_TESTING_REPORT.md',
            'output': 'FINAL_TESTING_REPORT.docx'
        },
        {
            'title': '✅ Checklist Completo de Testes',
            'input': 'COMPLETE_TESTING_CHECKLIST.md',
            'output': 'COMPLETE_TESTING_CHECKLIST.docx'
        },
        {
            'title': '🔧 Guia de Correção CORS',
            'input': 'CORS_FIX_GUIDE.md',
            'output': 'CORS_FIX_GUIDE.docx'
        },
        {
            'title': '🔍 Guia de Diagnóstico de Conexão',
            'input': 'DEBUG_CONNECTION.md',
            'output': 'DEBUG_CONNECTION.docx'
        },
        {
            'title': '🧪 Instruções de Teste',
            'input': 'TESTING_INSTRUCTIONS.md',
            'output': 'TESTING_INSTRUCTIONS.docx'
        },
        {
            'title': '📚 README - Gerador de Assinaturas',
            'input': 'README.md',
            'output': 'README.docx'
        }
    ]
    
    # Criar cada documento
    for doc_info in documents:
        try:
            create_document(
                doc_info['title'],
                doc_info['input'],
                doc_info['output']
            )
        except Exception as e:
            print(f'❌ Erro ao criar {doc_info["output"]}: {e}')
    
    print('\n🎉 Processo concluído!')
    print('\n📁 Documentos Word criados:')
    for doc_info in documents:
        if os.path.exists(doc_info['output']):
            print(f'  ✅ {doc_info["output"]}')

if __name__ == '__main__':
    main()
