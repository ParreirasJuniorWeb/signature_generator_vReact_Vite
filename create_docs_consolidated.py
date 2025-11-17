"""
Script para criar documento Word consolidado com toda documentação técnica
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def read_file(filepath):
    """Lê o conteúdo de um arquivo"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        return f'Arquivo não encontrado: {filepath}'
    except Exception as e:
        return f'Erro ao ler arquivo: {e}'

def add_formatted_content(doc, content):
    """Adiciona conteúdo formatado ao documento"""
    lines = content.split('\n')
    in_code_block = False
    code_lines = []
    
    for line in lines:
        # Detectar blocos de código
        if line.strip().startswith('```'):
            if in_code_block:
                # Fim do bloco de código
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph(code_text)
                    p.style = 'Intense Quote'
                code_lines = []
                in_code_block = False
            else:
                # Início do bloco de código
                in_code_block = True
            continue
        
        if in_code_block:
            code_lines.append(line)
            continue
        
        line_stripped = line.strip()
        
        if not line_stripped:
            doc.add_paragraph()
            continue
            
        # Títulos
        if line_stripped.startswith('# '):
            heading = doc.add_heading(line_stripped[2:], level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line_stripped.startswith('## '):
            doc.add_heading(line_stripped[3:], level=2)
        elif line_stripped.startswith('### '):
            doc.add_heading(line_stripped[4:], level=3)
        elif line_stripped.startswith('#### '):
            doc.add_heading(line_stripped[5:], level=4)
        
        # Listas com checkbox
        elif line_stripped.startswith('- [x] ') or line_stripped.startswith('- [ ] '):
            checkbox = '☑' if '[x]' in line_stripped else '☐'
            text = line_stripped.replace('- [x] ', '').replace('- [ ] ', '')
            doc.add_paragraph(f'{checkbox} {text}', style='List Bullet')
        
        # Listas normais
        elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
            doc.add_paragraph(line_stripped[2:], style='List Bullet')
        elif any(line_stripped.startswith(f'{i}. ') for i in range(1, 10)):
            dot_pos = line_stripped.find('. ')
            if dot_pos > 0:
                doc.add_paragraph(line_stripped[dot_pos+2:], style='List Number')
        
        # Texto normal
        else:
            doc.add_paragraph(line_stripped)

def create_consolidated_document():
    """Cria documento consolidado com toda documentação técnica"""
    print('🚀 Criando documento consolidado da documentação técnica...\n')
    
    doc = Document()
    
    # Capa
    title = doc.add_heading('📚 DOCUMENTAÇÃO TÉCNICA COMPLETA', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    subtitle = doc.add_heading('Gerador de Assinaturas de E-mail', level=1)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    info = doc.add_paragraph('Versão: 2.0.0')
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    date = doc.add_paragraph('Data: 16 de Novembro de 2025')
    date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    status = doc.add_paragraph('Status: ✅ DOCUMENTAÇÃO OFICIAL')
    status.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    status.runs[0].bold = True
    status.runs[0].font.color.rgb = RGBColor(0, 128, 0)
    
    doc.add_page_break()
    
    # Sumário
    doc.add_heading('📋 SUMÁRIO', level=1)
    doc.add_paragraph()
    
    toc_items = [
        '1. Resumo Executivo (SUMMARY)',
        '2. Arquitetura do Sistema (ARCHITECTURE)',
        '3. Documentação da API REST (API)',
        '4. Guia de Migração (MIGRATION_GUIDE)',
        '5. Histórico de Mudanças (CHANGELOG)',
        '6. Conclusão do Projeto (PROJECT_COMPLETION)'
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
        p.runs[0].bold = True
    
    doc.add_page_break()
    
    # Documentos
    documents = [
        {
            'title': '1. 📊 RESUMO EXECUTIVO',
            'file': 'docs/SUMMARY.md'
        },
        {
            'title': '2. 🏗️ ARQUITETURA DO SISTEMA',
            'file': 'docs/ARCHITECTURE.md'
        },
        {
            'title': '3. 📖 DOCUMENTAÇÃO DA API REST',
            'file': 'docs/API.md'
        },
        {
            'title': '4. 🔄 GUIA DE MIGRAÇÃO',
            'file': 'docs/MIGRATION_GUIDE.md'
        },
        {
            'title': '5. 📝 HISTÓRICO DE MUDANÇAS',
            'file': 'docs/CHANGELOG.md'
        },
        {
            'title': '6. ✅ CONCLUSÃO DO PROJETO',
            'file': 'docs/PROJECT_COMPLETION.md'
        }
    ]
    
    for doc_info in documents:
        print(f'Adicionando: {doc_info["title"]}')
        
        # Título da seção
        section_title = doc.add_heading(doc_info['title'], level=1)
        section_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()
        
        # Conteúdo
        content = read_file(doc_info['file'])
        add_formatted_content(doc, content)
        
        # Quebra de página (exceto no último)
        if doc_info != documents[-1]:
            doc.add_page_break()
    
    # Salvar
    output_file = 'DOCUMENTACAO_TECNICA_COMPLETA.docx'
    doc.save(output_file)
    print(f'\n✅ {output_file} criado com sucesso!')
    print(f'📁 Localização: {os.path.abspath(output_file)}')
    
    return output_file

if __name__ == '__main__':
    create_consolidated_document()
    print('\n🎉 Documento técnico consolidado criado com sucesso!')
